import os
import math
from datetime import datetime
from tkinter import filedialog, messagebox
from docx.shared import Inches

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _tight_paragraph(p, space_after=2):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)

def _safe_str(x, default=""):
    if x is None:
        return default
    s = str(x).strip()
    return s if s else default

def _parse_date_to_ddmmyyyy(date_str: str) -> str:
    date_str = _safe_str(date_str, "")
    for fmt in ("%m/%d/%y", "%m/%d/%Y", "%Y-%m-%d", "%d/%m/%Y", "%d/%m/%y"):
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime("%d%m%Y")
        except Exception:
            pass
    return datetime.today().strftime("%d%m%Y")


def _format_date_for_doc(date_str: str) -> str:
    date_str = _safe_str(date_str, "")
    for fmt in ("%m/%d/%y", "%m/%d/%Y", "%Y-%m-%d", "%d/%m/%Y", "%d/%m/%y"):
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime("%b %d, %Y")
        except Exception:
            pass
    return date_str or datetime.today().strftime("%b %d, %Y")


def _get_num_ribs(plan_data) -> str:
    choice = _safe_str(plan_data.get("anatomy.thoracic_count_choice"), "12")
    other = _safe_str(plan_data.get("anatomy.thoracic_count_other"), "")
    if choice.lower() == "other" and other:
        return other
    return choice


def _get_num_lumbar(plan_data) -> str:
    return _safe_str(plan_data.get("anatomy.lumbar_count"), "5")

def _get_levels_line(plan_data) -> str:
    sel = plan_data.get("level_selection", {}) or {}
    uiv = _safe_str(sel.get("uiv"), "—")
    liv = _safe_str(sel.get("liv"), "—")
    if uiv != "—" and liv != "—":
        return f"{uiv}-{liv}"
    return "—"


def _apply_base_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(12)


def _add_underlined_label_paragraph(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    run1 = p.add_run(label)
    run1.underline = True
    p.add_run(f" {text}")
    return p

def _build_rodfather_doc_lines(plan_data: dict):
    rs = (plan_data or {}).get("rod_selection", {}) or {}
    rf = rs.get("rod_father", {}) or {}
    if not rf:
        return []

    out = ["Rod Father:"]
    for side in ("left", "right"):
        info = rf.get(side)
        if not isinstance(info, dict):
            continue

        mode = _safe_str(info.get("mode"), "").lower()
        material = _safe_str(info.get("material"), "—")
        typ = _safe_str(info.get("type"), "—")

        req = info.get("required_length_mm")
        src_id = _safe_str(info.get("source_offcut_id"), "—")
        src_len = info.get("source_length_mm")

        left_len = info.get("leftover_length_mm")
        left_id = info.get("leftover_offcut_id")

        if mode == "exact":
            if req is not None:
                out.append(f"{side.capitalize()}: exact match {material}, {typ}, {float(req):.0f} mm, reserved {src_id}")
            else:
                out.append(f"{side.capitalize()}: exact match {material}, {typ}, reserved {src_id}")

        elif mode == "cut":
            req_txt = f"{float(req):.0f} mm" if req is not None else "—"
            src_txt = f"{float(src_len):.0f} mm" if src_len is not None else "—"
            if left_len is not None and float(left_len) > 0.5:
                out.append(f"{side.capitalize()}: cut {req_txt} from {src_txt} in Rod Father, created {float(left_len):.0f} mm offcut.")
            else:
                out.append(f"{side.capitalize()}: cut {req_txt} from {src_txt} in Rod Father, no leftover recorded.")

    return out if len(out) > 1 else []


def build_top_block(doc: Document, plan_data: dict):
    patient = plan_data.get("patient", {}) or {}

    patient_id = _safe_str(patient.get("id"), "UNKNOWN")
    age_years = _safe_str(patient.get("age_years"), "—")
    sex = _safe_str(patient.get("sex"), "—")
    months_post_menarchal = _safe_str(patient.get("months_post_menarchal"), "—")

    surgery_date_raw = _safe_str(patient.get("surgery_date"), "")
    date_line = _format_date_for_doc(surgery_date_raw)

    dx = _safe_str(patient.get("diagnosis"), "—")

    ribs = _get_num_ribs(plan_data)
    lumbar = _get_num_lumbar(plan_data)
    lld = _safe_str(plan_data.get("anatomy.lld"), "—")

    weight_kg = patient.get("weight_kg")
    weight_line = f"{weight_kg} kg" if weight_kg not in (None, "", "—") else "—"

    levels_line = _get_levels_line(plan_data)
    
    aim_val = patient.get("aim_text")

    if not aim_val:
        raw = patient.get("aim")
        if isinstance(raw, list):
            aim_val = ", ".join([str(x).strip() for x in raw if str(x).strip()])
        else:
            aim_val = str(raw).strip() if raw is not None else ""

    aim = _safe_str(aim_val, "—")

    # OP line
    p = doc.add_paragraph(f"OP {patient_id}")
    p.runs[0].bold = True
    _tight_paragraph(p)

    # Demographics
    p = doc.add_paragraph(f"{age_years} yo {sex}, {months_post_menarchal} months post menarchal.")
    _tight_paragraph(p)

    # Date
    p = doc.add_paragraph(date_line)
    _tight_paragraph(p)

    # Dx
    p = doc.add_paragraph(f"Dx: {dx}")
    _tight_paragraph(p)

    # Anatomy
    p = doc.add_paragraph(f"Anatomy: {ribs} ribs, {lumbar} lumbar vertebra, {lld}")
    _tight_paragraph(p)

    # Weight
    p = doc.add_paragraph(f"Weight: {weight_line}")
    _tight_paragraph(p)

    # Levels
    p = doc.add_paragraph(f"Levels - {levels_line}")
    _tight_paragraph(p)

    # Aim
    p = _add_underlined_label_paragraph(doc, "Aim:", aim)
    _tight_paragraph(p, space_after=6)

    # gap before the table
    gap = doc.add_paragraph("")
    _tight_paragraph(gap, space_after=4)

    build_anchor_table(doc, plan_data)

    # gap after table
    gap2 = doc.add_paragraph("")
    _tight_paragraph(gap2, space_after=4)

    screws_line = _summarize_screws_line(plan_data)
    p = _add_underlined_label_paragraph(doc, "Screws:", screws_line)
    _tight_paragraph(p)

    rods_line = _summarize_rods_line(plan_data)
    p = _add_underlined_label_paragraph(doc, "Rods:", rods_line)
    _tight_paragraph(p, space_after=4)

    rf_lines = _build_rodfather_doc_lines(plan_data)
    if rf_lines:
        for s in rf_lines:
            p2 = doc.add_paragraph(s)
            _tight_paragraph(p2, space_after=0)

        gap = doc.add_paragraph("")
        _tight_paragraph(gap, space_after=4)

    # Additional equipment
    add_additional_equipment_section_to_doc(doc, plan_data)

    # Positioning
    add_positioning_section_to_doc(doc, plan_data)

    add_infection_reduction_section_to_doc(doc, plan_data)

    add_blood_conservation_section_to_doc(doc, plan_data)

    add_correction_techniques_section_to_doc(doc, plan_data)
    add_post_op_pain_reduction_section_to_doc(doc, plan_data)
    add_post_op_destination_section_to_doc(doc, plan_data)

def export_docx_top_block(plan_data: dict, kind: str):
    """
    kind: "plan" or "op_note"
    """
    patient = plan_data.get("patient", {}) or {}
    patient_id = _safe_str(patient.get("id"), "UNKNOWN")
    surgery_date_raw = _safe_str(patient.get("surgery_date"), "")
    ddmmyyyy = _parse_date_to_ddmmyyyy(surgery_date_raw)

    suffix = "PLAN" if kind == "plan" else "OP NOTE"
    default_name = f"{patient_id} {ddmmyyyy} {suffix}.docx"

    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        initialfile=default_name,
        filetypes=[("Word Document", "*.docx")],
    )
    if not save_path:
        return

    doc = Document()
    _apply_base_style(doc)
    build_top_block(doc, plan_data)

    try:
        doc.save(save_path)
        messagebox.showinfo("Export Complete", f"Saved:\n{save_path}")
    except Exception as e:
        messagebox.showerror("Export Failed", str(e))

def _abbr_screw_type(s: str) -> str:
    s = _safe_str(s, "").lower()
    if "uniax" in s:
        return "UNI"
    if "poly" in s:
        return "POLY"
    if "mono" in s:
        return "MONO"
    if "cann" in s:
        return "CAN"
    return s.upper() if s else "SCREW"

def _format_anchor_side(side: dict) -> str:
    if not isinstance(side, dict):
        return "—"

    a = _safe_str(side.get("anchor_type"), "None")

    if a == "Screw":
        d = _safe_str(side.get("diameter_mm"), "")
        L = _safe_str(side.get("length_mm"), "")
        st = _abbr_screw_type(side.get("screw_type"))
        tap = bool(side.get("tap", False))

        if d and L:
            base = f"{d}/{L}/{st}"
        else:
            base = "SCREW"

        if tap:
            try:
                tap_d = float(d) - 1.0
                return base + f"\n(Tap {tap_d:.1f} mm)"
            except Exception:
                return base + "\n(Tap)"

        return base

    if a == "Tape":
        return "TAPE"

    if a == "Hook":
        return "HOOK"

    return "—"


def _get_rod_label(plan_data: dict, side: str) -> str:
    """
    Rod labels from Page 11:
      plan_data["rod_selection"]["left_rod"]
      plan_data["rod_selection"]["right_rod"]
    """
    rs = plan_data.get("rod_selection", {}) or {}
    key = "left_rod" if side == "left" else "right_rod"
    return _safe_str(rs.get(key), "")

def _format_rod_header(label: str) -> str:
    """
      'Cobalt Chrome 6.0 mm' -> '(6.0 CoCr)'
      'Cobalt Chrome 5 mm'   -> '(5 CoCr)'
      'Titanium 5.5 mm'     -> '(5.5 Titanium)'
    """
    s = _safe_str(label, "")
    if not s:
        return ""

    s = s.replace(" mm", "").strip()
    parts = s.split()

    diameter = None
    material_parts = []

    for part in parts:
        try:
            float(part)
            diameter = part
        except ValueError:
            material_parts.append(part)

    material = " ".join(material_parts)

    if "cobalt" in material.lower():
        material = "CoCr"

    if diameter:
        return f"({diameter} {material})"

    return f"({s})"


def _center_cell(cell):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_anchor_table(doc: Document, plan_data: dict):
    ap = plan_data.get("anchor_planning", {}) or {}
    levels = ap.get("levels", []) or []
    anchors = ap.get("anchors", {}) or {}

    if not levels:
        return 

    left_rod = _get_rod_label(plan_data, "left")
    right_rod = _get_rod_label(plan_data, "right")

    table = doc.add_table(rows=1, cols=3)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Level"
    
    hdr[1].text = "Left side"
    if left_rod:
        hdr[1].text = f"Left side\n{_format_rod_header(left_rod)}"

    hdr[2].text = "Right side"
    if right_rod:
        hdr[2].text = f"Right side\n{_format_rod_header(right_rod)}"

    _center_cell(hdr[1])
    _center_cell(hdr[2])

    # Header spacing 
    for c in hdr:
        for p in c.paragraphs:
            _tight_paragraph(p, space_after=0)

    # Data rows
    for lvl in levels:
        lvl = _safe_str(lvl, "")
        if not lvl:
            continue

        row = table.add_row().cells
        row[0].text = lvl

        lvl_data = anchors.get(lvl, {}) or {}
        left = _format_anchor_side(lvl_data.get("left", {}))
        right = _format_anchor_side(lvl_data.get("right", {}))

        row[1].text = left
        row[2].text = right

        _center_cell(row[1])
        _center_cell(row[2])

        for c in row:
            for p in c.paragraphs:
                _tight_paragraph(p, space_after=0)

def _collect_anchor_summary(plan_data: dict):
    ap = plan_data.get("anchor_planning", {}) or {}
    anchors = ap.get("anchors", {}) or {}

    screw_types = set()
    has_tape = False
    has_hook = False

    for lvl, lvl_data in anchors.items():
        if not isinstance(lvl_data, dict):
            continue
        for side_key in ("left", "right"):
            side = lvl_data.get(side_key, {}) or {}
            a = _safe_str(side.get("anchor_type"), "None")

            if a == "Screw":
                st = _safe_str(side.get("screw_type"), "")
                if st:
                    st_low = st.lower()
                    if "mono" in st_low:
                        screw_types.add("Monoblock")
                    elif "uniax" in st_low:
                        screw_types.add("Uniaxial")
                    elif "poly" in st_low:
                        screw_types.add("Polyaxial")
                    elif "cann" in st_low:
                        screw_types.add("Cannulated")
                    else:
                        screw_types.add(st.strip())
            elif a == "Tape":
                has_tape = True
            elif a == "Hook":
                has_hook = True

    return screw_types, has_tape, has_hook

def _summarize_screws_line(plan_data: dict) -> str:
    screw_types, has_tape, has_hook = _collect_anchor_summary(plan_data)

    parts = []

    parts.append("Stryker Xia")

    if has_tape:
        # change? clarify
        parts.append("Nile Sublaminar bands")

    # add screw types in a consistent order
    ordered = ["Monoblock", "Uniaxial", "Polyaxial", "Cannulated"]
    for t in ordered:
        if t in screw_types:
            parts.append(f"{t} screws")

    # optional, add only if any implants selected
    if screw_types or has_tape or has_hook:
        parts.append("Rod cut offs")

    return ", ".join(parts) if parts else "—"

def _parse_rod_label(label: str):
    s = _safe_str(label, "")
    if not s:
        return None, None

    s = s.replace(" mm", "").strip()
    parts = s.split()

    diameter = None
    material_parts = []
    for part in parts:
        try:
            float(part)
            diameter = part
        except ValueError:
            material_parts.append(part)

    material_raw = " ".join(material_parts).strip().lower()
    if "cobalt" in material_raw:
        material = "CoCr"
    elif "titan" in material_raw:
        material = "Ti-Alloy"
    else:
        material = " ".join(material_parts).strip() or "Rod"

    return material, diameter

def _summarize_rods_line(plan_data: dict) -> str:
    rs = plan_data.get("rod_selection", {}) or {}
    left = _safe_str(rs.get("left_rod"), "")
    right = _safe_str(rs.get("right_rod"), "")

    entries = []
    for lab in (left, right):
        mat, dia = _parse_rod_label(lab)
        if mat and dia:
            entries.append((mat, dia))

    if not entries:
        return "—"

    mat_to_dias = {}
    for mat, dia in entries:
        mat_to_dias.setdefault(mat, set()).add(dia)

    # if both materials exist and each has exactly one diameter and they match, combine
    if "CoCr" in mat_to_dias and "Ti-Alloy" in mat_to_dias:
        if len(mat_to_dias["CoCr"]) == 1 and len(mat_to_dias["Ti-Alloy"]) == 1:
            c_d = next(iter(mat_to_dias["CoCr"]))
            t_d = next(iter(mat_to_dias["Ti-Alloy"]))
            if c_d == t_d:
                return f"CoCr and Ti-Alloy rods {c_d} mm diameter."

    chunks = []
    for mat in ["CoCr", "Ti-Alloy"]:
        if mat not in mat_to_dias:
            continue

        dias = sorted(mat_to_dias[mat], key=lambda x: float(x))
        if len(dias) == 1:
            d = dias[0]
            if mat == "CoCr" and ("Ti-Alloy" in mat_to_dias):
                chunks.append(f"CoCr {d} mm diameter")
            else:
                chunks.append(f"{mat} rods {d} mm diameter")
        else:
            chunks.append(f"{mat} rods multiple diameters")

    if len(chunks) == 1:
        return chunks[0] + "."
    return " and ".join(chunks) + "."

def _table_type_to_text(setup: dict) -> str:
    t = (setup.get("table_type") or "").strip()
    other = (setup.get("table_other_text") or "").strip()

    if t == "trios":
        return 'Mizuho OSI Trios Spinal Surgery Top ("Jackson Spinal Bed")'
    if t == "neuro":
        return "Neurosurgical Table"
    if t == "other":
        return other or "Other table"
    return "Operating table not specified"


def _cranial_device_to_text(setup: dict) -> str:
    dev = (setup.get("cranial_device") or "").strip()
    mapping = {
        "none": "None",
        "gwtongs": "Gardner Wells Tongs",
        "mayfield": "Mayfield Clamp",
        "halo": "Halo via Mayfield Adapter",
    }
    return mapping.get(dev, dev or "Not specified")


def _femoral_type_to_text(setup: dict) -> str:
    t = (setup.get("femoral_type") or "").strip()
    mapping = {
        "none": "None",
        "skeletal": "Skeletal",
        "boots": "Traction Boots",
        "skin": "Skin Traction",
    }
    return mapping.get(t, t or "Not specified")


def _distribution_to_text(setup: dict) -> str:
    d = (setup.get("femoral_distribution") or setup.get("traction_distribution") or "").strip()
    if d == "symmetric":
        return "symmetric"
    if d == "asymmetric":
        return "asymmetric"
    return "not specified"


def _safe_float(x, default=None):
    try:
        return float(x)
    except Exception:
        return default


def build_positioning_line(plan_data: dict) -> str:
    setup = (plan_data or {}).get("setup", {}) or {}

    table_txt = _table_type_to_text(setup)

    traction_on = bool(setup.get("traction_on", False))
    if not traction_on:
        return f"{table_txt}; No traction"

    cranial = _safe_float(setup.get("cranial_weight_lbs"), default=0)
    fem_l = _safe_float(setup.get("femoral_left_lbs"), default=0)
    fem_r = _safe_float(setup.get("femoral_right_lbs"), default=0)
    fem_total = fem_l + fem_r

    fem_type_txt = _femoral_type_to_text(setup)
    cranial_dev = _cranial_device_to_text(setup)

    def _fmt(v):
        return str(int(round(v))) if abs(v - round(v)) < 1e-6 else f"{v:.1f}"

    base = f"{table_txt}; Cranial-Femoral Traction - {_fmt(cranial)}lbs/{_fmt(fem_total)}lbs"

    extras = []
    if fem_type_txt not in ("Not specified", "None"):
        extras.append(fem_type_txt)
    if cranial_dev not in ("Not specified", "None"):
        extras.append(f"Cranial Device: {cranial_dev}")

    if extras:
        base += f" ({', '.join(extras)})"

    return base

 
def add_positioning_section_to_doc(doc, plan_data: dict):
    line = build_positioning_line(plan_data)
    p = _add_underlined_label_paragraph(doc, "Positioning:", line)
    _tight_paragraph(p, space_after=6)  

def add_infection_reduction_section_to_doc(doc, plan_data: dict):
    inf = (plan_data or {}).get("infection_reduction", {}) or {}

    p = _add_underlined_label_paragraph(doc, "Infection Reduction Strategies:", "")
    _tight_paragraph(p, space_after=2)

    lines = []

    # Antibiotics
    if inf.get("pre_incision_abx", False):
        lines.append("Ancef prior to incision with re-dose during procedure")

    # Betadine paint
    if inf.get("povidone_paint_implants", False):
        lines.append("Betadine painting of implants prior to final closure when we ask for final x-rays")

    # Vancomycin powder
    vanc_wound = inf.get("vanc_wound_500mg", False)
    vanc_allo = inf.get("vanc_allograft_500mg", False)
    if vanc_wound or vanc_allo:
        if vanc_wound and vanc_allo:
            lines.append("1g Vancomycin powder [500mg combined with allograft and 500mg for wound closure]")
        elif vanc_allo:
            lines.append("Vancomycin powder 500mg combined with allograft")
        else:
            lines.append("Vancomycin powder 500mg for wound closure")

    # Print section
    if not lines:
        p = doc.add_paragraph("Antibiotics - —")
        _tight_paragraph(p, space_after=6)
        return

    first = True
    for line in lines:
        if first:
            p = doc.add_paragraph(f"Antibiotics - {line}")
            first = False
        else:
            p = doc.add_paragraph(f"- {line}")
        _tight_paragraph(p, space_after=0)

    gap = doc.add_paragraph("")
    _tight_paragraph(gap, space_after=6)

def build_additional_equipment_line(plan_data: dict) -> str:
    eq = (plan_data or {}).get("additional_equipment", {}) or {}

    parts = []

    # Neurophysiology
    if bool(eq.get("neuro_on", False)):
        modes = eq.get("neuro_modalities", {}) or {}
        chosen = [k for k in ("SSEPs", "MEPs", "EMGs") if bool(modes.get(k, False))]
        if chosen:
            parts.append(f"Neurophysiology ({', '.join(chosen)})")
        else:
            parts.append("Neurophysiology")

        baseline = (eq.get("neuro_baseline") or "").strip().lower()
        if baseline in ("supine", "prone"):
            base_txt = "Supine" if baseline == "supine" else "Prone (before traction)"
            parts.append(f"Baseline: {base_txt}")

    # Imaging / devices
    if bool(eq.get("small_cassette_on", False)):
        parts.append("Digital AP X-ray (small cassette) 45 minutes after incision (level check)")

    if bool(eq.get("sonopet_on", False)):
        parts.append("Sonapet")

    # 7D navigation + items
    if bool(eq.get("nav7d_on", False)):
        nav_items = eq.get("nav7d_items", {}) or {}
        nav_map = [
            ("pointer_ball_tip", "7D pointer (ball tip)"),
            ("pedicle_probe_lumbar", "Lumbar pedicle probe"),
            ("pedicle_probe_sharp", "Pedicle probe (sharp)"),
            ("spine_reference_clamp", "Spine reference clamp"),
            ("flex_array_rod_connector", "Flex array (flex rod connector)"),
        ]
        chosen_nav = [label for key, label in nav_map if bool(nav_items.get(key, False))]
        if chosen_nav:
            parts.append(f"7D Navigation [{', '.join(chosen_nav)}]")
        else:
            parts.append("7D Navigation")

    # Derotation
    if bool(eq.get("suk_on", False)):
        parts.append("De-rotation set (SUK DVR)")

    # Post-rod imaging
    if bool(eq.get("long_radiographs_on", False)):
        parts.append("3-foot AP and lateral radiographs after rods inserted")

    return ", ".join(parts) if parts else "—"


def add_additional_equipment_section_to_doc(doc, plan_data: dict):
    line = build_additional_equipment_line(plan_data)
    p = _add_underlined_label_paragraph(doc, "Additional equipment:", line)
    _tight_paragraph(p, space_after=6)


def _compute_infiltration_cocktail_lines(weight_kg: float):
    # Bupi 0.5% = 5 mg/ml, 1 mg/kg
    # Lido 1% = 10 mg/ml, 2 mg/kg
    # Epi = 5 mcg/ml of total volume, clamped to 1000 mcg in UI
    # NS = 200 ml
    bupi_mg = 1.0 * weight_kg
    lido_mg = 2.0 * weight_kg
    bupi_ml = bupi_mg / 5.0
    lido_ml = lido_mg / 10.0
    total_ml = bupi_ml + lido_ml + 200.0

    epi_total_mcg = 5.0 * total_ml
    if epi_total_mcg > 1000.0:
        epi_total_mcg = 1000.0

    # With your clamp, this is effectively always 1 ampule when enabled
    epi_ampules = int(math.ceil(epi_total_mcg / 1000.0)) if epi_total_mcg > 0 else 0

    return [
        f"Cocktail for this case ({weight_kg:.1f} kg):",
        f"{bupi_ml:.1f} ml of Bupivacaine 0.5%",
        f"{lido_ml:.1f} ml of Lidocaine 1%",
        f"{epi_ampules} ampule of epinephrine ({int(epi_total_mcg):d} mcg)",
        "200 ml of Normal Saline",
    ]

def _indent_paragraph(p, left_inches=0.35):
    p.paragraph_format.left_indent = Inches(left_inches)

def add_blood_conservation_section_to_doc(doc, plan_data: dict):
    blood = (plan_data or {}).get("blood_conservation", {}) or {}

    # Underlined section header
    p = _add_underlined_label_paragraph(doc, "Blood conservation strategies:", "")
    _tight_paragraph(p, space_after=2)

    n = 1  

    # 1) Infiltration strategy (numbered)
    if bool(blood.get("infiltration_on", False)):
        w = _safe_float(blood.get("infiltration_weight_kg"), default=None)
        if w is None:
            w = _safe_float((plan_data or {}).get("patient", {}).get("weight_kg"), default=None)

        if w is not None:
            # compute values for the mg/ml parentheticals
            bupi_mg = 1.0 * w
            lido_mg = 2.0 * w
            bupi_ml = bupi_mg / 5.0
            lido_ml = lido_mg / 10.0

            total_ml = bupi_ml + lido_ml + 200.0
            epi_total_mcg = 5.0 * total_ml
            if epi_total_mcg > 1000.0:
                epi_total_mcg = 1000.0
            epi_ampules = int(math.ceil(epi_total_mcg / 1000.0)) if epi_total_mcg > 0 else 0

            protocol = (
                f"Bupivacaine 0.5% (5mg/ml = 0.5%) (Astra Zeneca) 1 mg/Kg "
                f"({bupi_mg:.0f} mg = {bupi_ml:.0f} ml). "
                f"Lidocaine 1% (1% = 10mg/ml) 2 mg/Kg "
                f"({lido_mg:.0f} mg = {lido_ml:.0f} ml). "
                f"Epinephrine 5 mcg/ml of the total volume (1 ampule in 200 ml NS). "
                f"Add normal saline to a total volume of 100 ml/10 cm of the wound length."
            )

            p1 = doc.add_paragraph(f"{n}. {protocol}")
            _tight_paragraph(p1, space_after=2)
            n += 1

            # Cocktail header 
            cocktail_hdr = doc.add_paragraph()
            r = cocktail_hdr.add_run(f"Cocktail for this case ({w:.1f} kg):")
            r.underline = True
            _indent_paragraph(cocktail_hdr, left_inches=0.35)
            _tight_paragraph(cocktail_hdr, space_after=0)

            # Cocktail lines
            lines = [
                f"{bupi_ml:.0f} ml of Bupivacaine 0.5%",
                f"{lido_ml:.0f} ml of Lidocaine 1%",
                f"{epi_ampules} ampules of epinephrine ({int(epi_total_mcg):d} mcg)",
                "200 ml of Normal Saline",
            ]
            for s in lines:
                lp = doc.add_paragraph(s)
                _indent_paragraph(lp, left_inches=0.35)
                _tight_paragraph(lp, space_after=0)

            extra_notes = (blood.get("infiltration_notes") or "").strip()
            if extra_notes:
                np = doc.add_paragraph(extra_notes)
                _indent_paragraph(np, left_inches=0.35)
                _tight_paragraph(np, space_after=2)
            else:
                gap = doc.add_paragraph("")
                _tight_paragraph(gap, space_after=2)

    # 2) TXA
    if bool(blood.get("txa_on", True)):
        bolus = (blood.get("txa_bolus_mg_per_kg") or "").strip()
        inf = (blood.get("txa_infusion_mg_per_kg_hr") or "").strip()
        if bolus or inf:
            parts = []
            if bolus:
                parts.append(f"bolus {bolus} mg/kg")
            if inf:
                parts.append(f"infusion {inf} mg/kg/hr")
            txt = "Tranexamic Acid" + (f" ({', '.join(parts)})" if parts else "")
        else:
            txt = "Tranexamic Acid"

        p = doc.add_paragraph(f"{n}. {txt}")
        _tight_paragraph(p, space_after=0)
        n += 1

    # 3) CellSaver 
    if bool(blood.get("cell_saver_on", True)):
        p = doc.add_paragraph(f"{n}. CellSaver")
        _tight_paragraph(p, space_after=0)
        n += 1

    # 4) Floseal 
    if bool(blood.get("floseal_on", True)):
        raw_boxes = blood.get("floseal_boxes", None)

        boxes_int = None
        try:
            boxes_int = int(raw_boxes)
        except Exception:
            boxes_int = None

        loc = (blood.get("floseal_location") or "").strip()
        loc_txt = "In room" if loc == "in_room" else "Open" if loc == "open" else ""

        if boxes_int is not None and boxes_int > 0:
            txt = f"Floseal - {boxes_int} boxes"
        elif loc_txt:
            txt = f"Floseal - {loc_txt}"
        else:
            txt = "Floseal"

        p = doc.add_paragraph(f"{n}. {txt}")
        _tight_paragraph(p, space_after=0)
        n += 1

    gap = doc.add_paragraph("")
    _tight_paragraph(gap, space_after=6)

def _add_underlined_section_title(doc, title: str):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.underline = True
    _tight_paragraph(p, space_after=0)
    return p

def add_correction_techniques_section_to_doc(doc, plan_data):
    _add_underlined_section_title(doc, "Correction techniques:")
    p = doc.add_paragraph("Coming soon")
    _tight_paragraph(p, space_after=6)


def add_post_op_pain_reduction_section_to_doc(doc, plan_data):
    _add_underlined_section_title(doc, "Post-operative pain reduction techniques:")

    pain = plan_data.get("pain_reduction", {}) or {}
    pathway = (pain.get("pathway") or "").strip()
    notes = (pain.get("notes") or "").strip()

    if pathway == "intrathecal_morphine":
        line = "Intra-thecal morphine"
    elif pathway == "methadone":
        line = "Methadone"
    else:
        line = "—"

    p = doc.add_paragraph(line)
    _tight_paragraph(p, space_after=0)

    if notes:
        p2 = doc.add_paragraph(f"Notes: {notes}")
        _tight_paragraph(p2, space_after=6)
    else:
        _tight_paragraph(p, space_after=6)


def add_post_op_destination_section_to_doc(doc, plan_data):
    _add_underlined_section_title(doc, "Post-op Destination:")

    rec = plan_data.get("post_op_recovery", {}) or {}
    dest = (rec.get("destination") or "").strip()
    notes = (rec.get("notes") or "").strip()

    dest_map = {
        "5A_constant_obs": "5A Constant Observation Unit",
        "PICU": "Pediatric Intensive Care Unit (PICU)",
        "ICU_overnight": "Overnight Intensive Care Unit",
    }
    line = dest_map.get(dest, "—")

    p = doc.add_paragraph(line)
    _tight_paragraph(p, space_after=0)

    if notes:
        p2 = doc.add_paragraph(f"Notes: {notes}")
        _tight_paragraph(p2, space_after=6)
    else:
        _tight_paragraph(p, space_after=6)

def _build_postop_pain_text(plan_data: dict) -> str:
    pain = (plan_data or {}).get("pain_reduction", {}) or {}
    pathway = (pain.get("pathway") or "").strip()
    notes = (pain.get("notes") or "").strip()

    if pathway == "intrathecal_morphine":
        line = "Intra-thecal morphine"
    elif pathway == "methadone":
        line = "Methadone"
    else:
        line = "—"

    out = f"Post-operative pain reduction: {line}"
    if notes:
        out += f"\nNotes: {notes}"
    return out

# def _build_anchors_rods_text(plan_data: dict) -> str:
#     screws = _safe_str(_summarize_screws_line(plan_data), "—")
#     rods = _safe_str(_summarize_rods_line(plan_data), "—")
#     return f"Implants:\n- Screws: {screws}\n- Rods: {rods}"

def _build_anchors_rods_text(plan_data: dict) -> str:
    screws = _safe_str(_summarize_screws_line(plan_data), "—")
    rods = _safe_str(_summarize_rods_line(plan_data), "—")
    rf = _build_rodfather_summary_text(plan_data)
    return f"Implants:\n- Screws: {screws}\n- Rods: {rods}\n{rf}"


def _build_rodfather_summary_text(plan_data: dict) -> str:
    rs = (plan_data or {}).get("rod_selection", {}) or {}
    rf = rs.get("rod_father", {}) or {}

    if not rf:
        return "Rod Father: not used"

    lines = ["Rod Father:"]
    for side in ("left", "right"):
        info = rf.get(side)
        if not isinstance(info, dict):
            continue

        mode = (info.get("mode") or "").strip().lower()
        material = _safe_str(info.get("material"), "—")
        typ = _safe_str(info.get("type"), "—")

        req = info.get("required_length_mm")
        src_id = _safe_str(info.get("source_offcut_id"), "—")
        src_len = info.get("source_length_mm")

        left_len = info.get("leftover_length_mm")
        left_id = info.get("leftover_offcut_id")

        if mode == "exact":
            if req is not None:
                lines.append(f"- {side.capitalize()}: exact match {material}, {typ}, {float(req):.0f} mm, reserved")
            else:
                lines.append(f"- {side.capitalize()}: exact match {material}, {typ}, reserved")
        elif mode == "cut":
            req_txt = f"{float(req):.0f} mm" if req is not None else "—"
            src_txt = f"{float(src_len):.0f} mm" if src_len is not None else "—"
            if left_len is not None and float(left_len) > 0.5:
                lines.append(
                    f"- {side.capitalize()}: cut {req_txt} from {src_txt} in Rod Father, created {float(left_len):.0f} mm offcut."
                )
            else:
                lines.append(
                    f"- {side.capitalize()}: cut {req_txt} from {src_txt} in Rod Father, no leftover recorded."
                )
        else:
            lines.append(f"- {side.capitalize()}: used (details unavailable).")

    return "\n".join(lines)
